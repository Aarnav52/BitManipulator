import io
import json
import re
import time
import logging
from pathlib import Path
from typing import Optional, Tuple, List
from datetime import datetime

import pdfplumber
import fitz  # PyMuPDF
from google import genai
from google.genai import types
from docx import Document as DocxDocument

from core.config import get_settings
from models.resume import (
    ParsedResume, WorkExperience, Education, ExtractedSkill,
    Certification, Project, FileFormat, ProficiencyLevel,
)

logger = logging.getLogger(__name__)
settings = get_settings()

# ── Constants ──────────────────────────────────────────────────────────────────

SKIP_HEADERS = {
    "resume", "curriculum vitae", "cv", "profile", "summary",
    "objective", "contact", "about me", "about", "personal details",
    "personal information", "biodata", "bio-data", "candidate",
    "applicant", "name", "unknown candidate", "ai parsing failed",
    "professional summary", "career objective", "professional profile",
    "introduction", "overview", "bio",
    # Confidential / staffing headers
    "confidential", "confidential talent profile", "talent profile",
    "executive profile", "staffing director", "operations lead",
    # Section labels
    "core technical skills", "technical skills", "skills", "languages",
    "work experience", "experience", "education", "certifications",
    "projects", "references", "management", "soft skills",
    "management & soft skills", "core competencies", "competencies",
    "achievements", "awards", "publications", "interests",
    "hobbies", "volunteer",
}

# Words that indicate a job title line, not a name
TITLE_KEYWORDS = {
    "director", "manager", "lead", "engineer", "developer", "analyst",
    "consultant", "executive", "officer", "president", "vp", "head",
    "specialist", "coordinator", "architect", "designer", "staffing",
    "operations", "technical", "senior", "junior", "principal", "staff",
}

# ── Prompt ─────────────────────────────────────────────────────────────────────

EXTRACTION_SYSTEM_PROMPT = """You are an expert resume parser. Extract structured information into JSON.
Return ONLY valid JSON. No markdown backticks. No extra text.

Structure:
{
  "full_name": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin_url": "",
  "github_url": "",
  "portfolio_url": "",
  "summary": "",
  "total_years_experience": 0,
  "work_experience": [
    {
      "company": "",
      "role": "",
      "start_date": "",
      "end_date": "",
      "duration_months": 0,
      "responsibilities": [],
      "technologies": []
    }
  ],
  "education": [
    {
      "institution": "",
      "degree": "",
      "field": "",
      "year": null,
      "gpa": null
    }
  ],
  "skills": [
    {
      "raw_name": "",
      "proficiency": null,
      "years_experience": 0,
      "context_snippet": ""
    }
  ],
  "certifications": [
    {
      "name": "",
      "issuer": "",
      "year": null
    }
  ],
  "projects": [
    {
      "name": "",
      "description": "",
      "technologies": []
    }
  ]
}

Rules:
- Extract ALL skills mentioned anywhere in the resume
- Estimate total_years_experience from work history
- Use null for missing numeric fields, empty string for missing text
- proficiency values: beginner | familiar | intermediate | advanced | expert | null
- For full_name: extract the CANDIDATE'S actual human name only.
  IGNORE section labels like "Resume", "CV", "Curriculum Vitae", "Profile", "Summary", "Objective", "Contact", "About Me".
  A real name is typically 2-4 words near the top of the document, written in Title Case, containing ONLY letters (no digits, no special characters, no punctuation).
  Examples of valid names: "John Smith", "Priya Sharma", "Carlos de la Vega".
  If you cannot confidently identify a real person's name, return empty string "".
- Return ONLY the JSON, nothing else"""

EXTRACTION_USER_PROMPT = "Parse this resume:\n\n{raw_text}"


# ── JSON Cleaner ───────────────────────────────────────────────────────────────

def _clean_llm_json(raw: str) -> str:
    raw = raw.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    start = raw.find('{')
    end = raw.rfind('}')
    if start != -1 and end != -1:
        return raw[start:end + 1]
    return raw


# ── Confidence Calculator ──────────────────────────────────────────────────────

def _estimate_confidence(structured: dict) -> float:
    """Calculate confidence from extracted data — never trust LLM to self-report this."""
    score = 0.0
    name = structured.get("full_name", "")
    if name and name not in ("", "Unknown Candidate", "AI Parsing Failed"):
        score += 0.2
    if structured.get("email"):
        score += 0.1
    if structured.get("work_experience"):
        score += 0.3
    if structured.get("education"):
        score += 0.15
    if structured.get("skills"):
        score += 0.25
    return round(min(score, 1.0), 2)


# ── Name Extractor ─────────────────────────────────────────────────────────────

def _normalize_name_line(line: str) -> str:
    """
    Strip smart/curly quotes, asterisks, and other decoration from a line,
    then collapse whitespace. Keeps hyphens and apostrophes inside words.
    e.g. 'JONATHAN "JON" PRITCHARD'  →  'JONATHAN JON PRITCHARD'
         '*** CONFIDENTIAL ***'       →  'CONFIDENTIAL'
    """
    # Remove smart quotes, curly quotes, regular quotes
    line = re.sub(r'[\u201c\u201d\u2018\u2019"\']', '', line)
    # Remove decorative characters: asterisks, stars, pipes, dashes used as borders
    line = re.sub(r'[*|►▪•]', '', line)
    # Collapse whitespace
    line = re.sub(r'\s+', ' ', line).strip()
    return line


def _extract_name_from_text(text: str) -> str:
    """
    Heuristic name extractor. Scans the first 40 lines, skips headers/titles,
    and returns the first line that looks like a real person's name.
    Handles: Title Case, ALL CAPS, initials, nicknames in quotes,
             hyphenated names, smart quotes.
    """
    lines = text.splitlines()[:40]
    print(f"🔍 Scanning {len(lines)} lines for name...")

    for i, raw_line in enumerate(lines):
        line = _normalize_name_line(raw_line.strip())

        # Skip blank or too short/long after normalization
        if not line or len(line) < 3 or len(line) > 70:
            continue

        # Skip known section headers
        normalized = line.lower().strip(".:/-– \t")
        if normalized in SKIP_HEADERS:
            print(f"   Line {i}: '{line}' → skipped (known header)")
            continue

        # Skip lines that contain any title keyword — these are job title lines
        # e.g. "EXECUTIVE STAFFING DIRECTOR & OPERATIONS LEAD"
        words_lower = set(normalized.split())
        if words_lower & TITLE_KEYWORDS:
            print(f"   Line {i}: '{line}' → skipped (job title keywords)")
            continue

        # Skip lines with email / URLs
        if re.search(r'[@]', line):
            continue
        if re.search(r'(https?://|www\.)', line, re.I):
            continue

        # Skip lines with 3+ consecutive digits (phone, zip, date)
        if re.search(r'\d{3,}', line):
            continue

        # Skip lines with non-name special characters
        if re.search(r'[|•►▪✓\[\]<>{}#$%^&*_+=~`&]', line):
            continue

        # Skip address lines: "City, ST" or "City, Country"
        if re.search(r',\s*[A-Z][a-zA-Z]{1,}', line) and ',' in line:
            continue

        # Skip purely lowercase lines
        if line.islower():
            continue

        words = line.split()

        # Names are 2–5 words
        if not (2 <= len(words) <= 5):
            continue

        # Normalize ALL-CAPS to Title Case for validation
        display = line
        if line.isupper():
            display = line.title()
            words = display.split()

        def is_name_word(w: str) -> bool:
            # Initials: "R.", "R.K."
            if re.match(r'^[A-Z](\.[A-Z])*\.?$', w):
                return True
            # Single capital letter initial (no dot)
            if re.match(r'^[A-Z]$', w):
                return True
            # Regular name word: capital start, letters/hyphens/apostrophes
            if re.match(r"^[A-Z][a-zA-Z\-']{1,}$", w):
                return True
            return False

        if all(is_name_word(w) for w in words):
            print(f"   ✅ Name found at line {i}: '{display}'")
            return display

        print(f"   Line {i}: '{line}' → no word match (words={words})")

    print("   ❌ No name found via heuristics")
    return ""


# ── LLM Call ───────────────────────────────────────────────────────────────────

def structure_with_llm(raw_text: str) -> Tuple[dict, List[str]]:
    warnings: List[str] = []

    api_key = getattr(settings, 'google_api_key', None)

    if not api_key or "your_" in str(api_key):
        print("❌ No Google API Key found in settings!")
        warnings.append("Missing API Key - using regex fallback")
        return _regex_fallback(raw_text), warnings

    # Truncate to avoid token limits
    words = raw_text.split()
    if len(words) > 6000:
        raw_text = " ".join(words[:6000])
        warnings.append("Resume truncated to 6000 words")

    try:
        client = genai.Client(api_key=api_key)

        print("🚀 Sending to Gemini...")
        response = client.models.generate_content(
            model=settings.llm_model,
            contents=EXTRACTION_USER_PROMPT.format(raw_text=raw_text),
            config=types.GenerateContentConfig(
                system_instruction=EXTRACTION_SYSTEM_PROMPT,
            ),
        )

        if not response.text:
            print("⚠️ Gemini returned empty response")
            warnings.append("Empty AI response - using regex fallback")
            return _regex_fallback(raw_text), warnings

        cleaned = _clean_llm_json(response.text)
        data = json.loads(cleaned)

        # ── Post-process: validate and fix the extracted name ──────────────────
        extracted_name = data.get("full_name", "").strip()
        print(f"🧠 Gemini returned full_name: '{extracted_name}'")

        name_is_bad = (
            not extracted_name
            or extracted_name.lower().strip(".:/-– ") in SKIP_HEADERS
            or len(extracted_name.split()) < 2
            or re.search(r'[\d@]', extracted_name)
        )

        if name_is_bad:
            print(f"⚠️ Name failed validation, switching to heuristic extraction...")
            warnings.append(f"LLM name '{extracted_name}' failed validation, using heuristic extraction")
            heuristic_name = _extract_name_from_text(raw_text)
            print(f"   Heuristic result: '{heuristic_name}'")
            data["full_name"] = heuristic_name or "Unknown Candidate"

        print(f"✅ Final name: '{data['full_name']}' | Skills: {len(data.get('skills', []))}")
        return data, warnings

    except json.JSONDecodeError as e:
        print(f"❌ JSON parse error: {e}")
        warnings.append(f"LLM returned invalid JSON: {e}")
        return _regex_fallback(raw_text), warnings
    except Exception as e:
        print(f"❌ Gemini Error: {e}")
        warnings.append(str(e))
        return _regex_fallback(raw_text), warnings


# ── Regex Fallback ─────────────────────────────────────────────────────────────

def _regex_fallback(text: str) -> dict:
    print("⚠️ Using Regex Fallback...")
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[a-z]{2,}', text, re.I)
    phone_match = re.search(r'(\+?\d[\d\s\-(). ]{7,}\d)', text)
    linkedin_match = re.search(r'linkedin\.com/in/[\w\-]+', text, re.I)
    github_match = re.search(r'github\.com/[\w\-]+', text, re.I)

    # Use the robust heuristic extractor instead of blind first-line grab
    name = _extract_name_from_text(text)

    # Try to extract skills section
    skills_raw = []
    skill_section = re.search(
        r'(?:skills?|technical skills?)[:\s]*(.*?)(?:\n\n|\Z)', text, re.I | re.S
    )
    if skill_section:
        skills_raw = [
            s.strip(" *-|,")
            for s in re.split(r'[,\n|*]', skill_section.group(1))
            if 1 < len(s.strip()) < 60
        ][:30]

    return {
        "full_name": name or "Unknown Candidate",
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0).strip() if phone_match else "",
        "location": "",
        "linkedin_url": f"https://{linkedin_match.group(0)}" if linkedin_match else "",
        "github_url": f"https://{github_match.group(0)}" if github_match else "",
        "portfolio_url": "",
        "summary": "",
        "total_years_experience": None,
        "work_experience": [],
        "education": [],
        "skills": [{"raw_name": s, "proficiency": None, "years_experience": 0, "context_snippet": ""} for s in skills_raw],
        "certifications": [],
        "projects": [],
    }


# ── Safe Helpers ───────────────────────────────────────────────────────────────

def _safe_proficiency(value) -> Optional[ProficiencyLevel]:
    if not value:
        return None
    try:
        return ProficiencyLevel(str(value).lower())
    except ValueError:
        return None


def _safe_float(value, default=None):
    try:
        return float(value) if value is not None else default
    except (TypeError, ValueError):
        return default


def _safe_int(value, default=None):
    try:
        return int(value) if value is not None else default
    except (TypeError, ValueError):
        return default


# ── Model Builder ──────────────────────────────────────────────────────────────

def build_parsed_resume(structured, source_format, source_filename, raw_text, warnings, parse_start) -> ParsedResume:

    # Skills
    skills_list = []
    for s in structured.get("skills", []):
        if isinstance(s, dict) and s.get("raw_name", "").strip():
            skills_list.append(ExtractedSkill(
                raw_name=str(s["raw_name"]).strip(),
                proficiency=_safe_proficiency(s.get("proficiency")),
                years_experience=_safe_float(s.get("years_experience")),
                context_snippet=str(s.get("context_snippet", ""))[:120],
            ))
        elif isinstance(s, str) and s.strip():
            skills_list.append(ExtractedSkill(raw_name=s.strip()))

    # Work experience
    work_list = []
    for w in structured.get("work_experience", []):
        if isinstance(w, dict):
            work_list.append(WorkExperience(
                company=w.get("company") or "",
                role=w.get("role") or "",
                start_date=w.get("start_date") or "",
                end_date=w.get("end_date") or "",
                duration_months=_safe_int(w.get("duration_months")),
                responsibilities=w.get("responsibilities") or [],
                technologies=w.get("technologies") or [],
            ))

    # Education
    education_list = []
    for e in structured.get("education", []):
        if isinstance(e, dict):
            education_list.append(Education(
                institution=e.get("institution") or "",
                degree=e.get("degree") or "",
                field=e.get("field") or "",
                year=_safe_int(e.get("year")),
                gpa=_safe_float(e.get("gpa")),
            ))

    # Certifications
    cert_list = []
    for c in structured.get("certifications", []):
        if isinstance(c, dict):
            cert_list.append(Certification(
                name=c.get("name") or "",
                issuer=c.get("issuer") or "",
                year=_safe_int(c.get("year")),
            ))

    # Projects
    project_list = []
    for p in structured.get("projects", []):
        if isinstance(p, dict):
            project_list.append(Project(
                name=p.get("name") or "",
                description=p.get("description") or "",
                technologies=p.get("technologies") or [],
            ))

    confidence = _estimate_confidence(structured)

    return ParsedResume(
        full_name=structured.get("full_name") or "Unknown Candidate",
        email=structured.get("email") or None,
        phone=structured.get("phone") or None,
        location=structured.get("location") or None,
        linkedin_url=structured.get("linkedin_url") or None,
        github_url=structured.get("github_url") or None,
        portfolio_url=structured.get("portfolio_url") or None,
        summary=structured.get("summary") or "",
        total_years_experience=_safe_float(structured.get("total_years_experience")),
        work_experience=work_list,
        education=education_list,
        skills=skills_list,
        certifications=cert_list,
        projects=project_list,
        source_format=source_format,
        source_filename=source_filename,
        raw_text_length=len(raw_text),
        parsing_confidence=float(confidence),
        parse_warnings=warnings,
        parsed_at=datetime.utcnow(),
    )


# ── File Validator ─────────────────────────────────────────────────────────────

def _validate_file(file_bytes: bytes, filename: str) -> Tuple[bool, str]:
    """
    Checks if the file is actually a real resume document.
    Returns (is_valid, error_message).
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        if not file_bytes[:4] == b"%PDF":
            return False, "File has .pdf extension but is not a real PDF."
        return True, ""

    if ext in (".docx", ".doc"):
        if not file_bytes[:2] == b"PK":
            return False, "File has .docx extension but is not a valid Word document."

        import zipfile
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                names = z.namelist()
                if "word/document.xml" not in names:
                    content_hint = ", ".join(names[:5])
                    return False, (
                        f"This is not a resume DOCX — it appears to be an Office theme or template file. "
                        f"Please upload your actual resume. (Contains: {content_hint})"
                    )
        except zipfile.BadZipFile:
            return False, "File is corrupted and cannot be opened."
        return True, ""

    return True, ""


# ── DOCX Extractor ─────────────────────────────────────────────────────────────

def _extract_docx_text(file_bytes: bytes) -> Tuple[str, List[str]]:
    """
    Robust DOCX text extraction — tries multiple strategies.
    """
    warnings: List[str] = []

    # Strategy 1: python-docx
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        if len(text.strip()) > 50:
            return text, warnings
        warnings.append("python-docx returned sparse text, trying XML fallback")
    except Exception as e:
        warnings.append(f"python-docx failed: {e}, trying XML fallback")

    # Strategy 2: Raw XML extraction
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
                texts = [
                    node.text
                    for node in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                    if node.text
                ]
                text = " ".join(texts)
                if len(text.strip()) > 50:
                    return text, warnings
                warnings.append("XML fallback also returned sparse text")
    except Exception as e:
        warnings.append(f"XML fallback failed: {e}")

    return "", warnings


# ── Parsing Agent ──────────────────────────────────────────────────────────────

class ParsingAgent:
    def parse_file(self, file_bytes: bytes, filename: str) -> ParsedResume:
        t0 = time.perf_counter()
        warnings: List[str] = []

        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            fmt = FileFormat.PDF
        elif ext in [".docx", ".doc"]:
            fmt = FileFormat.DOCX
        else:
            fmt = FileFormat.TXT

        # ── Validate file before trying to parse ──
        is_valid, validation_error = _validate_file(file_bytes, filename)
        if not is_valid:
            print(f"❌ Invalid file: {validation_error}")
            warnings.append(validation_error)
            error_data = _regex_fallback("")
            error_data["full_name"] = "Invalid File"
            return build_parsed_resume(error_data, fmt, filename, "", warnings, t0)

        # ── Extract text ──
        raw_text = ""
        try:
            if fmt == FileFormat.PDF:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    pages = [p.extract_text() or "" for p in pdf.pages]
                    raw_text = "\n\n".join(pages)

                if len(raw_text.strip()) < 100:
                    warnings.append("pdfplumber returned sparse text, trying PyMuPDF")
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    raw_text = "\n\n".join(page.get_text("text") for page in doc)
                    doc.close()

            elif fmt == FileFormat.DOCX:
                raw_text, docx_warnings = _extract_docx_text(file_bytes)
                warnings.extend(docx_warnings)

            else:
                raw_text = file_bytes.decode("utf-8", errors="ignore")

        except Exception as e:
            warnings.append(f"Text extraction error: {e}")
            raw_text = ""

        if len(raw_text.strip()) < 20:
            warnings.append("Very little text extracted from file")
            return build_parsed_resume(
                _regex_fallback(raw_text), fmt, filename, raw_text, warnings, t0
            )

        structured, llm_warnings = structure_with_llm(raw_text)
        warnings.extend(llm_warnings)
        return build_parsed_resume(structured, fmt, filename, raw_text, warnings, t0)


parsing_agent = ParsingAgent()